# -*- coding: utf-8 -*-
"""
医考帮学习备份工具 - Streamlit 网页版（云端适配）
兼容医考帮新版 ENC1 加密题目文件与旧版 JSON 题目文件。
"""

import hashlib
import io
import json
import os
import re
import tempfile
import time
import zipfile
from urllib.parse import urlparse

import pandas as pd
import requests
import streamlit as st
from tqdm import tqdm


st.set_page_config(page_title="医考帮备份工具", page_icon="📚", layout="centered")

# ============================================================
# 页面访问密码
# ============================================================

APP_ACCESS_PASSWORD = os.environ.get("APP_ACCESS_PASSWORD", "268369")


def check_password():
    """显示密码输入框，验证通过才让进。"""
    if "app_unlocked" not in st.session_state:
        st.session_state.app_unlocked = False
    if st.session_state.app_unlocked:
        return True

    st.title("🔒 请输入访问密码")
    password_input = st.text_input("密码", type="password", placeholder="请输入密码")
    if st.button("登录", type="primary"):
        if password_input == APP_ACCESS_PASSWORD:
            st.session_state.app_unlocked = True
            st.rerun()
        else:
            st.error("❌ 密码错误，请重试")
    return False


if not check_password():
    st.stop()

# ============================================================
# 导出模块
# ============================================================


def fix_encoding(text):
    if not text:
        return text
    try:
        return text.encode("latin1").decode("utf-8")
    except Exception:
        return text


def _download_image(url, save_path):
    try:
        r = requests.get(url, stream=True, timeout=10)
        r.raise_for_status()
        with open(save_path, "wb") as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)
    except Exception:
        pass


def _parse_options(option_value):
    if not option_value:
        return []
    if isinstance(option_value, list):
        return option_value
    if isinstance(option_value, str):
        try:
            parsed = json.loads(option_value)
            return parsed if isinstance(parsed, list) else []
        except Exception:
            return []
    return []


def export_exam(json_data, output_path_base, format="excel", download_images=False):
    """导出考试数据，所有文件保存在临时目录下。"""
    output_dir = os.path.dirname(output_path_base)
    img_dir = os.path.join(output_dir, "images")

    rows = []
    for q in tqdm(json_data, desc="解析题目", unit="题"):
        opts = _parse_options(q.get("option", []))
        opt_str = "\n".join(
            [f"{o.get('key', '')}: {fix_encoding(o.get('title', ''))}" for o in opts if isinstance(o, dict)]
        )
        title = fix_encoding(q.get("title", ""))
        explain = fix_encoding(q.get("explain", ""))

        if download_images:
            all_text = f"{title}{explain}{opt_str}"
            img_urls = re.findall(r"(https?://[^\s\"'<>]+\.(?:jpg|jpeg|png|gif))", all_text, re.IGNORECASE)
            if img_urls:
                os.makedirs(img_dir, exist_ok=True)
                for img_url in img_urls:
                    filename = os.path.basename(urlparse(img_url).path) or "image.jpg"
                    save_path = os.path.join(img_dir, filename)
                    if not os.path.exists(save_path):
                        _download_image(img_url, save_path)

        rows.append(
            {
                "题号": q.get("number", ""),
                "题目": title,
                "选项": opt_str,
                "答案": q.get("answer", ""),
                "解析": explain,
            }
        )

    if format == "excel":
        output_path = f"{output_path_base}.xlsx"
        pd.DataFrame(rows).to_excel(output_path, index=False)
    elif format == "word":
        output_path = f"{output_path_base}.docx"
        _export_word(rows, output_path)
    elif format == "pdf":
        output_path = f"{output_path_base}.pdf"
        _export_pdf(rows, output_path)


def _export_word(rows, output_path):
    from docx import Document

    doc = Document()
    doc.add_heading("医考帮题库导出", 0)
    for row in tqdm(rows, desc="生成 Word", unit="题"):
        doc.add_heading(f"第 {row['题号']} 题", level=2)
        doc.add_paragraph(f"【题目】{row['题目']}")
        doc.add_paragraph(f"【选项】\n{row['选项']}")
        p = doc.add_paragraph()
        p.add_run(f"【正确答案】{row['答案']}").bold = True
        doc.add_paragraph(f"【解析】{row['解析']}")
        doc.add_paragraph("─" * 30)
    doc.save(output_path)


def _export_pdf(rows, output_path):
    import textwrap

    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    font_name = "Helvetica"
    font_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
    ]
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("CJK", fp))
                font_name = "CJK"
                break
            except Exception:
                continue

    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    margin = 50
    line_height = 16
    y = height - margin

    def new_page():
        nonlocal y
        c.showPage()
        c.setFont(font_name, 10)
        y = height - margin

    c.setFont(font_name, 10)

    for row in tqdm(rows, desc="生成 PDF", unit="题"):
        lines = [
            f"第 {row['题号']} 题",
            f"题目：{row['题目']}",
        ]
        for opt_line in row["选项"].split("\n"):
            lines.append(f"  {opt_line}")
        lines.append(f"正确答案：{row['答案']}")
        lines.append(f"解析：{row['解析']}")
        lines.append("")

        for line in lines:
            wrapped = textwrap.wrap(line, width=60) if line else [""]
            for wline in wrapped:
                if y < margin + line_height:
                    new_page()
                c.drawString(margin, y, wline)
                y -= line_height

    c.save()


# ============================================================
# 登录与 API 模块
# ============================================================

BASE_URL = "https://ykb-online-exam.yikaobang.com.cn/api"
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "application/json, text/plain, */*",
    "Content-Type": "application/x-www-form-urlencoded;charset=UTF-8",
    "Client-Type": "web",
    "Origin": "https://ykb-online-exam.yikaobang.com.cn",
    "Referer": "https://ykb-online-exam.yikaobang.com.cn/",
}

DOWNLOAD_HEADERS = {
    "User-Agent": HEADERS["User-Agent"],
    "Accept": "application/json, text/plain, */*",
    "Referer": "https://ykb-online-exam.yikaobang.com.cn/",
}

# 医考帮网页端当前 ENC1 题目文件密钥版本。
ENC1_KEY_STORE = {
    1: "9ef5eb2e388a0f149cb85b070b013a217201b52260ba896926a4771d613e347e",
}


def md5(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def _request_with_retry(method, url, **kwargs):
    last_exc = None
    for attempt in range(3):
        try:
            response = requests.request(method, url, **kwargs)
            if response.status_code in (502, 503, 504) and attempt < 2:
                time.sleep(0.8 * (attempt + 1))
                continue
            return response
        except requests.RequestException as exc:
            last_exc = exc
            if attempt < 2:
                time.sleep(0.8 * (attempt + 1))
                continue
    raise RuntimeError(f"网络请求失败：{last_exc}")


def _json_response(response, context):
    if response.status_code != 200:
        raise RuntimeError(f"{context}失败：HTTP {response.status_code}")
    if not response.content:
        raise RuntimeError(f"{context}失败：服务器返回空内容")
    try:
        return response.json()
    except ValueError:
        preview = response.text[:120].replace("\n", " ")
        raise RuntimeError(f"{context}失败：返回内容不是 JSON，前缀：{preview}")


def login(mobile, password):
    url = f"{BASE_URL}/user/login"
    data = {"mobile": mobile, "password": md5(password)}
    try:
        r = _request_with_retry("post", url, data=data, headers=HEADERS, timeout=(10, 30))
        result = _json_response(r, "登录")
        if result.get("code") == 200:
            token = result["data"]["token"]
            secret = result["data"]["secret"]
            hospitals = result["data"].get("hospital", [])
            hospital_id = hospitals[0]["id"] if hospitals else ""
            return token, secret, str(hospital_id)
        st.error(result.get("message", "登录失败，请检查账号密码"))
        return None
    except Exception as e:
        st.error(f"登录失败：{e}")
        return None


def get_exam_list(token, secret, hospital_id):
    url = f"{BASE_URL}/question/getUserExam"
    data = {"token": token, "secret": secret, "hospital_id": hospital_id}
    try:
        r = _request_with_retry("post", url, data=data, headers=HEADERS, timeout=(10, 40))
        result = _json_response(r, "获取考试列表")
        if result.get("code") == 200:
            return result.get("data", [])
        st.error(result.get("message", "获取考试列表失败"))
        return []
    except Exception as e:
        st.error(f"获取考试列表失败：{e}")
        return []


def _normalize_question_url(url):
    if not url:
        return url
    # 医考帮返回的 OSS 链接仍可能是 http；新版环境下 http 偶发 502 空响应。
    if url.startswith("http://ksxt-image-1.oss-cn-beijing.aliyuncs.com/"):
        return "https://" + url[len("http://") :]
    return url


def _decode_plain_json(content):
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _decrypt_enc1(content):
    try:
        from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    except ImportError as exc:
        raise RuntimeError("缺少 cryptography 依赖，请在 requirements.txt 中加入 cryptography 后重新部署") from exc

    if len(content) < 40:
        raise RuntimeError("ENC1 文件过短，无法解析")
    if content[:4] != b"ENC1":
        raise RuntimeError("不是有效的 ENC1 文件")

    fmt = content[4]
    key_version = content[5]
    if fmt != 1:
        raise RuntimeError(f"暂不支持的 ENC1 格式版本：{fmt}")

    key_hex = ENC1_KEY_STORE.get(key_version)
    if not key_hex:
        raise RuntimeError(f"未知的 ENC1 密钥版本：{key_version}")

    iv = content[8:20]
    expected_size = int.from_bytes(content[20:24], "big")
    ciphertext = content[24:-16]
    tag = content[-16:]

    key = bytes.fromhex(key_hex)
    plaintext = AESGCM(key).decrypt(iv, ciphertext + tag, None)
    if len(plaintext) != expected_size:
        raise RuntimeError("ENC1 解密后大小校验失败")
    return plaintext.decode("utf-8")


def _parse_question_payload(content):
    if not content:
        raise RuntimeError("题目文件为空")

    stripped = content.lstrip()
    if stripped.startswith((b"[", b"{")):
        return json.loads(_decode_plain_json(stripped))
    if content.startswith(b"ENC1"):
        return json.loads(_decrypt_enc1(content))

    preview = content[:20].hex()
    raise RuntimeError(f"无法识别题目文件格式，文件头：{preview}")


def download_questions(question_file_url):
    url = _normalize_question_url(question_file_url)
    try:
        r = _request_with_retry("get", url, headers=DOWNLOAD_HEADERS, timeout=(10, 60), allow_redirects=True)
        if r.status_code != 200:
            raise RuntimeError(f"HTTP {r.status_code}")
        return _parse_question_payload(r.content)
    except Exception as e:
        raise RuntimeError(f"题目数据下载或解析失败：{e}") from e


# ============================================================
# 网页界面
# ============================================================


def main():
    st.title("📚 医考帮个人学习备份工具")
    st.caption("基于 ykb-exporter 改造 · 网页版")
    st.divider()

    if "auth" not in st.session_state:
        st.session_state.auth = None
    if "exams" not in st.session_state:
        st.session_state.exams = []
    if "ykb_logged_in" not in st.session_state:
        st.session_state.ykb_logged_in = False

    with st.sidebar:
        st.header("🔐 登录账号")
        mobile = st.text_input("手机号", placeholder="请输入手机号")
        password = st.text_input("密码", placeholder="请输入密码", type="password")

        if st.button("登录", type="primary", use_container_width=True):
            if not mobile or not password:
                st.error("请完整填写手机号和密码")
            else:
                with st.spinner("正在登录，请稍候..."):
                    auth = login(mobile, password)
                    if auth:
                        st.session_state.auth = auth
                        st.session_state.ykb_logged_in = True
                        with st.spinner("正在获取考试列表..."):
                            st.session_state.exams = get_exam_list(auth[0], auth[1], auth[2])
                        st.success("✅ 登录成功！")
                        st.rerun()

        if st.session_state.ykb_logged_in:
            st.info(f"当前已登录：{mobile}")
            if st.button("退出登录"):
                st.session_state.auth = None
                st.session_state.exams = []
                st.session_state.ykb_logged_in = False
                st.rerun()

    if not st.session_state.ykb_logged_in:
        st.info("👈 请先在左侧侧边栏登录您的医考帮账号")
        return

    exams = st.session_state.exams
    if not exams:
        st.warning("未获取到任何考试，请确认账号权限或重新登录")
        return

    st.subheader("📋 选择要导出的考试")

    status_map = {0: "未开始", 1: "进行中", 2: "已结束"}
    exam_options = []
    for e in exams:
        status = status_map.get(e.get("status", 0), "未知")
        title = e.get("title", "未命名")
        exam_options.append(f"[{status}] {title}")

    selected_index = st.selectbox(
        "考试列表（请点击选择）",
        options=range(len(exam_options)),
        format_func=lambda x: exam_options[x],
    )

    st.subheader("📄 选择导出格式")
    formats = st.multiselect(
        "支持同时导出多种格式",
        options=["excel", "word", "pdf"],
        default=["excel", "word", "pdf"],
        format_func=lambda x: {"excel": "📊 Excel (.xlsx)", "word": "📝 Word (.docx)", "pdf": "📕 PDF (.pdf)"}[x],
    )

    if not formats:
        st.warning("请至少选择一种导出格式")

    if st.button("🚀 开始导出", type="primary", use_container_width=True, disabled=(not formats)):
        selected_exam = exams[selected_index]
        title = selected_exam.get("title", "exam_export")
        question_url = selected_exam.get("question_file")

        if not question_url:
            st.error("该考试没有题目文件链接，无法导出")
            st.stop()

        status_placeholder = st.status("⏳ 正在处理，请耐心等待...", expanded=True)

        try:
            status_placeholder.write("📥 正在下载题目数据...")
            json_data = download_questions(question_url)
            if not json_data:
                st.error("题目数据为空，无法导出")
                status_placeholder.update(label="导出失败", state="error")
                st.stop()
            status_placeholder.write(f"✅ 下载完成，共 {len(json_data)} 道题目")

            with tempfile.TemporaryDirectory() as tmpdir:
                safe_title = re.sub(r'[\\/:*?"<>|]', "_", title)
                base_path = os.path.join(tmpdir, safe_title)

                status_placeholder.write("🔄 正在转换格式...")
                for fmt in formats:
                    status_placeholder.write(f"   - 生成 {fmt} 文件...")
                    export_exam(json_data, base_path, format=fmt, download_images=False)

                status_placeholder.write("📦 正在打包文件...")
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    for f in os.listdir(tmpdir):
                        file_path = os.path.join(tmpdir, f)
                        if os.path.isfile(file_path) and f.endswith((".xlsx", ".docx", ".pdf")):
                            zf.write(file_path, f)
                zip_buffer.seek(0)

                status_placeholder.update(label="✅ 导出完成！", state="complete")

                st.success(f"🎉 成功导出 {len(formats)} 个文件！点击下方按钮下载：")
                st.download_button(
                    label="📥 点击下载 ZIP 压缩包",
                    data=zip_buffer,
                    file_name=f"{safe_title}_export.zip",
                    mime="application/zip",
                    type="primary",
                    use_container_width=True,
                )

        except Exception as e:
            status_placeholder.update(label="❌ 导出出错", state="error")
            st.error(f"处理过程中发生错误：{e}")
            st.stop()

    st.divider()
    st.caption("💡 提示：所有数据仅在临时目录处理，不会长期保存在服务器")


if __name__ == "__main__":
    main()
